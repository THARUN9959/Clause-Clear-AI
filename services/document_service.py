"""Document parsing service — extracts text from PDF, DOCX, TXT, and images."""

import os
import logging

from config import Config

logger = logging.getLogger(__name__)

try:
    from PyPDF2 import PdfReader
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed — PDF uploads not supported.")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX uploads not supported.")

try:
    import pytesseract
    from PIL import Image
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    logger.warning("pytesseract/Pillow not installed — OCR not available.")


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in Config.ALLOWED_EXTENSIONS


def extract_text(filepath: str) -> str:
    """
    Extract text from an uploaded document. All file handles are opened in
    'with' context managers so Windows file locks are released before any
    os.remove() call by the caller.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        text = _extract_pdf_text(filepath)
        if not text.strip() and _OCR_AVAILABLE:
            logger.info("PDF text layer empty; attempting OCR: %s", filepath)
            text = _extract_pdf_ocr(filepath)
        return text

    if ext == ".txt":
        return _extract_txt(filepath)

    if ext == ".docx":
        return _extract_docx(filepath)

    if ext in {".png", ".jpg", ".jpeg"}:
        return _extract_image_ocr(filepath)

    return ""


def _extract_pdf_text(filepath: str) -> str:
    if not _PDF_AVAILABLE:
        return ""
    parts = []
    try:
        with open(filepath, "rb") as fh:
            reader = PdfReader(fh)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
    except Exception as exc:
        logger.error("PyPDF2 error for %s: %s", filepath, exc)
    return "\n".join(parts)


def _extract_pdf_ocr(filepath: str) -> str:
    try:
        from pdf2image import convert_from_path
    except ImportError:
        logger.warning("pdf2image not installed — cannot OCR scanned PDFs.")
        return ""
    parts = []
    try:
        pages = convert_from_path(filepath, dpi=200)
        for page_img in pages:
            text = pytesseract.image_to_string(page_img)
            if text:
                parts.append(text)
    except Exception as exc:
        logger.error("OCR PDF error for %s: %s", filepath, exc)
    return "\n".join(parts)


def _extract_txt(filepath: str) -> str:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()
    except Exception as exc:
        logger.error("TXT read error for %s: %s", filepath, exc)
        return ""


def _extract_docx(filepath: str) -> str:
    if not _DOCX_AVAILABLE:
        return ""
    try:
        with open(filepath, "rb") as fh:
            doc = DocxDocument(fh)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
        return "\n".join(parts)
    except Exception as exc:
        logger.error("DOCX extraction error for %s: %s", filepath, exc)
        return ""


def _extract_image_ocr(filepath: str) -> str:
    if not _OCR_AVAILABLE:
        return ""
    try:
        with Image.open(filepath) as img:
            return pytesseract.image_to_string(img)
    except Exception as exc:
        logger.error("Image OCR error for %s: %s", filepath, exc)
        return ""
